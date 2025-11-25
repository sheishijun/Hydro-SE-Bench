"""
汇总报告生成功能
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

try:
    import pandas as pd
    import numpy as np
except ImportError:
    pd = None  # type: ignore[assignment]
    np = None  # type: ignore[assignment]

from .batch_evaluate import identify_model_columns
from .benchmark import Benchmark


def _generate_word_report(md_lines: list[str], output_dir: Path) -> None:
    """将 Markdown 内容转换为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
    except ImportError:
        print("  ⚠ python-docx 未安装，跳过 Word 报告生成")
        print("    提示: 运行 pip install python-docx 可生成 Word 格式报告")
        return
    
    try:
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'  # 中文字体
        font.size = Pt(10.5)
        
        i = 0
        while i < len(md_lines):
            line = md_lines[i].strip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 图片处理（Markdown 格式：![alt](path)）
            if line.startswith('![') and '](' in line:
                # 提取图片路径
                try:
                    img_path = line.split('](')[1].split(')')[0]
                    img_file = output_dir / img_path
                    if img_file.exists() and img_file.is_file():
                        doc.add_picture(str(img_file), width=Inches(6))
                        doc.add_paragraph()  # 添加空行
                except Exception:
                    pass  # 如果图片不存在，跳过
                i += 1
                continue
            
            # 标题处理
            if line.startswith('# '):
                # 一级标题
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                # 二级标题
                p = doc.add_heading(line[3:], level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                # 三级标题
                p = doc.add_heading(line[4:], level=3)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('---'):
                # 分隔线
                doc.add_paragraph('─' * 50)
            elif line.startswith('|'):
                # 表格
                table_lines = []
                while i < len(md_lines) and md_lines[i].strip().startswith('|'):
                    table_lines.append(md_lines[i].strip())
                    i += 1
                i -= 1  # 回退一步
                
                if len(table_lines) >= 2:
                    # 解析表头
                    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                    # 跳过分隔行
                    data_rows = table_lines[2:]
                    
                    # 创建表格
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # 添加表头
                    header_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        header_cells[j].text = header
                        header_cells[j].paragraphs[0].runs[0].font.bold = True
                    
                    # 添加数据行
                    for row_line in data_rows:
                        if row_line.strip().startswith('| ...'):
                            continue  # 跳过省略行
                        cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        if len(cells) == len(headers):
                            row_cells = table.add_row().cells
                            for j, cell_text in enumerate(cells):
                                row_cells[j].text = cell_text
            elif line.startswith('- ') or line.startswith('* '):
                # 列表项
                text = line[2:].strip()
                # 处理粗体
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                if text:  # 确保文本不为空
                    doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\d+\.\s+', line):
                # 有序列表（匹配任何数字开头的列表项）
                text = re.sub(r'^\d+\.\s+', '', line)
                # 处理粗体
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                if text:  # 确保文本不为空
                    doc.add_paragraph(text, style='List Number')
            else:
                # 普通段落
                # 处理粗体
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                if text:  # 确保文本不为空
                    doc.add_paragraph(text)
            
            i += 1
        
        # 保存 Word 文档
        word_file = output_dir / "detailed_analysis_report.docx"
        doc.save(str(word_file))
        print(f"✓ 深度分析报告（Word）已保存: {word_file}")
    except Exception as e:
        print(f"  ⚠ Word 报告生成失败: {e}")


def _generate_pdf_report(md_lines: list[str], output_dir: Path) -> None:
    """将 Markdown 内容转换为 PDF 文档"""
    try:
        import markdown
    except ImportError:
        print("  ⚠ markdown 未安装，跳过 PDF 报告生成")
        print("    提示: 运行 pip install markdown 可生成 PDF 格式报告")
        print("    或者: 使用 Word 文档在 Microsoft Word 中另存为 PDF")
        return
    
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
    except (ImportError, OSError) as e:
        print("  ⚠ weasyprint 未安装或系统依赖缺失，跳过 PDF 报告生成")
        print(f"    错误: {e}")
        print("    提示: weasyprint 在 Windows 上需要额外的系统库，安装较复杂")
        print("    建议: 使用 Word 文档在 Microsoft Word 中另存为 PDF")
        return
    
    try:
        # 将 Markdown 转换为 HTML
        md_content = "\n".join(md_lines)
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code']
        )
        
        # 添加样式
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    font-size: 24pt;
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    margin-top: 30px;
                }}
                h2 {{
                    font-size: 18pt;
                    color: #34495e;
                    margin-top: 25px;
                    border-bottom: 2px solid #ecf0f1;
                    padding-bottom: 5px;
                }}
                h3 {{
                    font-size: 14pt;
                    color: #7f8c8d;
                    margin-top: 20px;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                    font-size: 10pt;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #3498db;
                    color: white;
                    font-weight: bold;
                }}
                tr:nth-child(even) {{
                    background-color: #f9f9f9;
                }}
                ul, ol {{
                    margin: 10px 0;
                    padding-left: 30px;
                }}
                li {{
                    margin: 5px 0;
                }}
                p {{
                    margin: 10px 0;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # 生成 PDF
        pdf_file = output_dir / "detailed_analysis_report.pdf"
        font_config = FontConfiguration()
        HTML(string=styled_html).write_pdf(
            pdf_file,
            font_config=font_config
        )
        print(f"✓ 深度分析报告（PDF）已保存: {pdf_file}")
    except (OSError, Exception) as e:
        print(f"  ⚠ PDF 报告生成失败: {e}")
        print("    提示: weasyprint 在 Windows 上可能需要额外的系统库")
        print("    建议: 使用 Word 文档在 Microsoft Word 中另存为 PDF")


def create_summary_excel(summary: dict[str, Any], output_dir: Path, benchmark: Any = None) -> None:
    """
    创建汇总 Excel 报告，包含所有模型的对比和按类别/难度的统计。
    
    Args:
        summary: 评估汇总结果字典，包含 results、benchmark、total_questions 等字段
        output_dir: 输出目录路径
        benchmark: Benchmark 对象（可选，当前未使用但保留接口兼容性）
    """
    if pd is None:
        print("⚠ pandas 未安装，跳过汇总 Excel 报告生成")
        return
    
    # 模型对比表
    rows = []
    for result in summary["results"]:
        rows.append({
            "Model Name": result["model_name"],
            "Correct Count": result["total_score"],
            "Total": result["max_score"],
            "Accuracy": result["accuracy"],  # 使用数值，Excel会自动格式化为百分比
            "Incorrect": result["incorrect_count"],
            "Missing": result["missing_count"],
        })
    
    df = pd.DataFrame(rows)
    
    summary_file = output_dir / "models_comparison.xlsx"
    with pd.ExcelWriter(summary_file, engine="openpyxl") as writer:
        # 模型对比工作表
        df.to_excel(writer, sheet_name="Model Comparison", index=False)
        worksheet = writer.sheets["Model Comparison"]
        worksheet.column_dimensions["A"].width = 40  # Model Name
        worksheet.column_dimensions["B"].width = 12  # Correct Count
        worksheet.column_dimensions["C"].width = 12  # Total
        worksheet.column_dimensions["D"].width = 12  # Accuracy
        worksheet.column_dimensions["E"].width = 10  # Incorrect
        worksheet.column_dimensions["F"].width = 10  # Missing
        
        # 设置准确率列为百分比格式
        from openpyxl.styles import numbers
        for row in range(2, len(df) + 2):  # 从第2行开始（第1行是表头）
            cell = worksheet.cell(row, 4)  # D列是准确率
            cell.number_format = numbers.FORMAT_PERCENTAGE_00  # 百分比格式，保留2位小数
        
        # 添加汇总信息
        summary_row = len(df) + 3
        worksheet.cell(summary_row, 1, "Benchmark Info")
        worksheet.cell(summary_row + 1, 1, "Benchmark: hydrobench")
        worksheet.cell(summary_row + 2, 1, f"Total Questions: {summary['total_questions']}")
        worksheet.cell(summary_row + 3, 1, f"Models Evaluated: {summary['models_count']}")
        
        # 按类别统计（所有模型的对比）
        if any(result.get("category_stats") for result in summary["results"]):
            # 收集所有类别
            all_categories = set()
            for result in summary["results"]:
                if result.get("category_stats"):
                    all_categories.update(result["category_stats"].keys())
            
            if all_categories:
                category_rows = []
                for category in sorted(all_categories):
                    row = {"Category": category}
                    for result in summary["results"]:
                        model_name = result["model_name"]
                        stats = result.get("category_stats", {}).get(category, {})
                        if stats:
                            # 分开两列：得分和准确率
                            row[f"{model_name}_Score"] = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                            row[f"{model_name}_Accuracy"] = stats.get('accuracy', 0)  # 数值格式
                        else:
                            row[f"{model_name}_Score"] = "-"
                            row[f"{model_name}_Accuracy"] = None
                    category_rows.append(row)
                
                category_df = pd.DataFrame(category_rows)
                category_df.to_excel(writer, sheet_name="By Category Comparison", index=False)
                cat_worksheet = writer.sheets["By Category Comparison"]
                cat_worksheet.column_dimensions["A"].width = 30  # Category
                
                # 设置列宽和格式
                col_idx = 2
                for result in summary["results"]:
                    model_name = result["model_name"]
                    # 得分列
                    cat_worksheet.column_dimensions[chr(64 + col_idx)].width = 20
                    col_idx += 1
                    # 准确率列 - 设置为百分比格式
                    from openpyxl.styles import numbers
                    for row in range(2, len(category_df) + 2):
                        cell = cat_worksheet.cell(row, col_idx)
                        if cell.value is not None:
                            cell.number_format = numbers.FORMAT_PERCENTAGE_00
                    cat_worksheet.column_dimensions[chr(64 + col_idx)].width = 15
                    col_idx += 1
        
        # 按难度统计（所有模型的对比）
        if any(result.get("level_stats") for result in summary["results"]):
            # 收集所有难度
            all_levels = set()
            for result in summary["results"]:
                if result.get("level_stats"):
                    all_levels.update(result["level_stats"].keys())
            
            if all_levels:
                level_rows = []
                for level in sorted(all_levels):
                    row = {"Level": level}
                    for result in summary["results"]:
                        model_name = result["model_name"]
                        stats = result.get("level_stats", {}).get(level, {})
                        if stats:
                            # 分开两列：得分和准确率
                            row[f"{model_name}_Score"] = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                            row[f"{model_name}_Accuracy"] = stats.get('accuracy', 0)  # 数值格式
                        else:
                            row[f"{model_name}_Score"] = "-"
                            row[f"{model_name}_Accuracy"] = None
                    level_rows.append(row)
                
                level_df = pd.DataFrame(level_rows)
                level_df.to_excel(writer, sheet_name="By Level Comparison", index=False)
                level_worksheet = writer.sheets["By Level Comparison"]
                level_worksheet.column_dimensions["A"].width = 30  # Level
                
                # 设置列宽和格式
                col_idx = 2
                for result in summary["results"]:
                    model_name = result["model_name"]
                    # 得分列
                    level_worksheet.column_dimensions[chr(64 + col_idx)].width = 20
                    col_idx += 1
                    # 准确率列 - 设置为百分比格式
                    from openpyxl.styles import numbers
                    for row in range(2, len(level_df) + 2):
                        cell = level_worksheet.cell(row, col_idx)
                        if cell.value is not None:
                            cell.number_format = numbers.FORMAT_PERCENTAGE_00
                    level_worksheet.column_dimensions[chr(64 + col_idx)].width = 15
                    col_idx += 1
        
        # 按题型统计（所有模型的对比）
        if any(result.get("type_stats") for result in summary["results"]):
            # 收集所有题型
            all_types = set()
            for result in summary["results"]:
                if result.get("type_stats"):
                    all_types.update(result["type_stats"].keys())
            
            if all_types:
                type_rows = []
                for qtype in sorted(all_types):
                    row = {"Type": qtype}
                    for result in summary["results"]:
                        model_name = result["model_name"]
                        stats = result.get("type_stats", {}).get(qtype, {})
                        if stats:
                            # 分开两列：得分和准确率
                            row[f"{model_name}_Score"] = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                            row[f"{model_name}_Accuracy"] = stats.get('accuracy', 0)  # 数值格式
                        else:
                            row[f"{model_name}_Score"] = "-"
                            row[f"{model_name}_Accuracy"] = None
                    type_rows.append(row)
                
                type_df = pd.DataFrame(type_rows)
                type_df.to_excel(writer, sheet_name="By Type Comparison", index=False)
                type_worksheet = writer.sheets["By Type Comparison"]
                type_worksheet.column_dimensions["A"].width = 30  # Type
                
                # 设置列宽和格式
                col_idx = 2
                for result in summary["results"]:
                    model_name = result["model_name"]
                    # 得分列
                    type_worksheet.column_dimensions[chr(64 + col_idx)].width = 20
                    col_idx += 1
                    # 准确率列 - 设置为百分比格式
                    from openpyxl.styles import numbers
                    for row in range(2, len(type_df) + 2):
                        cell = type_worksheet.cell(row, col_idx)
                        if cell.value is not None:
                            cell.number_format = numbers.FORMAT_PERCENTAGE_00
                    type_worksheet.column_dimensions[chr(64 + col_idx)].width = 15
                    col_idx += 1
    
    print(f"✓ 模型对比汇总已保存: {summary_file}")


def validate_data_quality(df: pd.DataFrame, benchmark: Benchmark) -> dict[str, Any]:
    """
    验证数据质量。
    
    Args:
        df: 预测数据的 DataFrame
        benchmark: Benchmark 对象
    
    Returns:
        包含数据质量检查结果的字典
    """
    if pd is None:
        raise ImportError("pandas 未安装，无法进行数据质量检查")
    
    issues = []
    warnings = []
    
    # 检查 ID 列
    if "ID" not in df.columns:
        issues.append("缺少 ID 列")
    else:
        # 检查 ID 匹配度
        df_ids = set(df["ID"].astype(str))
        benchmark_ids = set(ex.id for ex in benchmark.examples)
        missing_ids = benchmark_ids - df_ids
        extra_ids = df_ids - benchmark_ids
        
        if missing_ids:
            warnings.append(f"测评集中有 {len(missing_ids)} 个 ID 在预测文件中缺失")
        if extra_ids:
            warnings.append(f"预测文件中有 {len(extra_ids)} 个 ID 不在测评集中")
    
    # 检查空值
    model_cols = identify_model_columns(df, verbose=False)
    for col in model_cols:
        null_count = df[col].isna().sum()
        if null_count > 0:
            warnings.append(f"模型列 '{col}' 有 {null_count} 个空值")
    
    return {
        "issues": issues,
        "warnings": warnings,
        "model_count": len(model_cols),
        "data_rows": len(df),
        "benchmark_size": len(benchmark.examples),
    }


def analyze_errors(summary: dict[str, Any], benchmark: Benchmark, output_dir: Path) -> dict[str, Any]:
    """
    深度错误分析。
    
    Args:
        summary: 评估汇总结果
        benchmark: Benchmark 对象
        output_dir: 输出目录，用于读取详细报告
    
    Returns:
        错误分析结果字典
    """
    error_analysis = {
        "common_errors": {},  # 题目ID -> 错误模型列表
        "hard_questions": [],  # 所有模型都错的题目
        "easy_questions": [],  # 所有模型都对的题目
        "model_errors": {},  # 模型 -> 错误题目列表
    }
    
    # 收集每道题的错误情况
    question_errors = {}  # question_id -> {correct_models: [], incorrect_models: []}
    
    for result in summary["results"]:
        model_name = result["model_name"]
        error_analysis["model_errors"][model_name] = []
        
        # 读取该模型的详细报告
        # 生成安全的文件夹名
        safe_name = model_name.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?", "_").replace('"', "_").replace("<", "_").replace(">", "_").replace("|", "_")
        model_dir = output_dir / safe_name
        report_file = model_dir / "score_report.json"
        if report_file.exists():
            with open(report_file, "r", encoding="utf-8") as f:
                report_data = json.load(f)
            
            for example in report_data.get("examples", []):
                qid = example["id"]
                is_correct = example["is_correct"]
                
                if qid not in question_errors:
                    question_errors[qid] = {"correct": [], "incorrect": []}
                
                if is_correct:
                    question_errors[qid]["correct"].append(model_name)
                else:
                    question_errors[qid]["incorrect"].append(model_name)
                    error_analysis["model_errors"][model_name].append(qid)
    
    # 分析题目难度
    total_models = len(summary["results"])
    for qid, stats in question_errors.items():
        correct_count = len(stats["correct"])
        incorrect_count = len(stats["incorrect"])
        
        if incorrect_count == total_models:
            error_analysis["hard_questions"].append(qid)
        elif correct_count == total_models:
            error_analysis["easy_questions"].append(qid)
        
        if incorrect_count > total_models * 0.5:  # 超过一半模型都错
            error_analysis["common_errors"][qid] = {
                "error_rate": incorrect_count / total_models,
                "incorrect_models": stats["incorrect"],
            }
    
    return error_analysis


def generate_analysis_report(
    summary: dict[str, Any],
    benchmark: Benchmark,
    output_dir: Path,
) -> dict[str, Any]:
    """
    生成深度分析报告（Markdown 格式）。
    
    Args:
        summary: 评估汇总结果
        benchmark: Benchmark 对象
        output_dir: 输出目录
    
    Returns:
        分析报告字典（JSON 格式）
    """
    if pd is None:
        raise ImportError("pandas 未安装，无法生成分析报告")
    
    best_model = summary["results"][0]
    worst_model = summary["results"][-1]
    avg_accuracy = sum(r["accuracy"] for r in summary["results"]) / len(summary["results"])
    std_accuracy = pd.Series([r["accuracy"] for r in summary["results"]]).std()
    accuracy_gap = best_model["accuracy"] - worst_model["accuracy"]
    
    # 生成 Markdown 报告
    md_lines = []
    md_lines.append("# 深度分析报告")
    md_lines.append("")
    md_lines.append(f"**生成时间**: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    
    # 1. 评估概览
    md_lines.append("## 📊 评估概览")
    md_lines.append("")
    md_lines.append(f"- **评估模型数**: {summary['models_count']} 个")
    md_lines.append(f"- **题目总数**: {summary['total_questions']} 道")
    md_lines.append(f"- **平均准确率**: {avg_accuracy:.2%}")
    md_lines.append(f"- **准确率标准差**: {std_accuracy:.4f}")
    md_lines.append("")
    
    # 2. 模型排名
    md_lines.append("## 🏆 模型排名")
    md_lines.append("")
    md_lines.append("| 排名 | 模型名称 | 得分 | 准确率 |")
    md_lines.append("|------|---------|------|--------|")
    for idx, result in enumerate(summary["results"], 1):
        score_str = f"{result['total_score']}/{result['max_score']}"
        accuracy_str = f"{result['accuracy']:.2%}"
        md_lines.append(f"| {idx} | {result['model_name']} | {score_str} | {accuracy_str} |")
    md_lines.append("")
    
    # 3. 关键指标
    md_lines.append("## 📈 关键指标")
    md_lines.append("")
    md_lines.append(f"### 最佳模型")
    md_lines.append(f"- **模型**: {best_model['model_name']}")
    md_lines.append(f"- **准确率**: {best_model['accuracy']:.2%}")
    md_lines.append(f"- **得分**: {best_model['total_score']}/{best_model['max_score']}")
    md_lines.append("")
    md_lines.append(f"### 最差模型")
    md_lines.append(f"- **模型**: {worst_model['model_name']}")
    md_lines.append(f"- **准确率**: {worst_model['accuracy']:.2%}")
    md_lines.append(f"- **得分**: {worst_model['total_score']}/{worst_model['max_score']}")
    md_lines.append("")
    md_lines.append(f"### 模型差距")
    md_lines.append(f"- **准确率差距**: {accuracy_gap:.2%}")
    md_lines.append("")
    
    # 4. 类别分析
    if summary["results"] and summary["results"][0].get("category_stats"):
        md_lines.append("## 📂 类别维度分析")
        md_lines.append("")
        all_categories = set()
        for result in summary["results"]:
            if result.get("category_stats"):
                all_categories.update(result["category_stats"].keys())
        
        for category in sorted(all_categories):
            md_lines.append(f"### {category}")
            md_lines.append("")
            md_lines.append("| 模型名称 | 准确率 | 得分 |")
            md_lines.append("|---------|--------|------|")
            
            best_acc = 0
            best_model_name = ""
            worst_acc = 1.0
            worst_model_name = ""
            
            category_results = []
            for result in summary["results"]:
                stats = result.get("category_stats", {}).get(category, {})
                if stats:
                    acc = stats.get("accuracy", 0)
                    score_str = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                    category_results.append((result["model_name"], acc, score_str))
                    if acc > best_acc:
                        best_acc = acc
                        best_model_name = result["model_name"]
                    if acc < worst_acc:
                        worst_acc = acc
                        worst_model_name = result["model_name"]
            
            # 按准确率排序
            category_results.sort(key=lambda x: x[1], reverse=True)
            for model_name, acc, score_str in category_results:
                md_lines.append(f"| {model_name} | {acc:.2%} | {score_str} |")
            
            md_lines.append("")
            md_lines.append(f"- **最佳**: {best_model_name} ({best_acc:.2%})")
            md_lines.append(f"- **最差**: {worst_model_name} ({worst_acc:.2%})")
            md_lines.append(f"- **差距**: {(best_acc - worst_acc):.2%}")
            md_lines.append("")
    
    # 5. 难度分析
    if any(r.get("level_stats") for r in summary["results"]):
        md_lines.append("## 🎯 难度维度分析")
        md_lines.append("")
        all_levels = set()
        for result in summary["results"]:
            if result.get("level_stats"):
                all_levels.update(result["level_stats"].keys())
        
        md_lines.append("| 难度级别 | 平均准确率 | 模型数 |")
        md_lines.append("|---------|-----------|--------|")
        
        for level in sorted(all_levels):
            level_accuracies = []
            for result in summary["results"]:
                stats = result.get("level_stats", {}).get(level, {})
                if stats:
                    level_accuracies.append(stats.get("accuracy", 0))
            
            if level_accuracies:
                avg_level_acc = sum(level_accuracies) / len(level_accuracies)
                md_lines.append(f"| {level} | {avg_level_acc:.2%} | {len(level_accuracies)} |")
        
        md_lines.append("")
    
    # 6. 建议
    md_lines.append("## 💡 关键建议")
    md_lines.append("")
    recommendations = []
    recommendations.append(f"✅ 最佳模型 **{best_model['model_name']}** 准确率为 {best_model['accuracy']:.2%}，建议作为基准模型")
    recommendations.append(f"⚠️ 模型间准确率差距为 {accuracy_gap:.2%}，建议分析表现差异原因")
    
    if avg_accuracy < 0.5:
        recommendations.append("⚠️ 平均准确率低于 50%，建议检查模型配置或数据质量")
    elif avg_accuracy > 0.9:
        recommendations.append("✅ 平均准确率较高，建议增加题目难度或扩大测评集规模")
    
    for i, rec in enumerate(recommendations, 1):
        md_lines.append(f"{i}. {rec}")
    md_lines.append("")
    
    # 8. 详细数据（链接到 JSON）
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 📎 相关文件")
    md_lines.append("")
    md_lines.append("- 完整数据（JSON 格式）: `detailed_analysis_report.json`")
    md_lines.append("- 模型对比汇总: `models_comparison.xlsx`")
    md_lines.append("- 各模型详细报告: `<模型名>/score_report.xlsx` 和 `<模型名>/score_report.json`")
    md_lines.append("- Word 格式报告: `detailed_analysis_report.docx`（如果安装了 python-docx）")
    md_lines.append("- PDF 格式报告: `detailed_analysis_report.pdf`（如果安装了 weasyprint）")
    md_lines.append("")
    
    # 保存 Markdown 报告
    md_file = output_dir / "detailed_analysis_report.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    
    print(f"✓ 深度分析报告（Markdown）已保存: {md_file}")
    
    # 尝试生成 Word 和 PDF 格式
    _generate_word_report(md_lines, output_dir)
    _generate_pdf_report(md_lines, output_dir)
    
    # 同时保存 JSON 格式作为数据备份
    report_json = {
        "summary": {
            "total_models": summary["models_count"],
            "total_questions": summary["total_questions"],
            "evaluation_date": pd.Timestamp.now().isoformat(),
            "average_accuracy": avg_accuracy,
            "std_accuracy": float(std_accuracy),
        },
        "model_rankings": [
            {
                "rank": idx,
                "model_name": result["model_name"],
                "accuracy": result["accuracy"],
                "score": f"{result['total_score']}/{result['max_score']}",
            }
            for idx, result in enumerate(summary["results"], 1)
        ],
        "key_metrics": {
            "best_model": {
                "name": best_model["model_name"],
                "accuracy": best_model["accuracy"],
                "score": f"{best_model['total_score']}/{best_model['max_score']}",
            },
            "worst_model": {
                "name": worst_model["model_name"],
                "accuracy": worst_model["accuracy"],
                "score": f"{worst_model['total_score']}/{worst_model['max_score']}",
            },
            "accuracy_gap": accuracy_gap,
        },
        "recommendations": recommendations,
    }
    
    json_file = output_dir / "detailed_analysis_report.json"
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(report_json, f, ensure_ascii=False, indent=2)
    
    print(f"✓ 深度分析报告（JSON 数据）已保存: {json_file}")
    
    return report_json

