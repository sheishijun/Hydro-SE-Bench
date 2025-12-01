"""
Summary report generation functionality
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

try:
    import pandas as pd
    import numpy as np
except ImportError:
    pd = None  # type: ignore[assignment]
    np = None  # type: ignore[assignment]

from .batch_evaluate import identify_model_columns
from .benchmark import Benchmark


def _generate_word_report(md_lines: list[str], output_dir: Path) -> None:
    """Convert Markdown content to Word document"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
    except ImportError:
        print("  ⚠ python-docx not installed, skipping Word report generation")
        print("    Tip: Run pip install python-docx to generate Word format reports")
        return
    
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'  # Chinese font
        font.size = Pt(10.5)
        
        i = 0
        while i < len(md_lines):
            line = md_lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Image processing (Markdown format: ![alt](path))
            if line.startswith('![') and '](' in line:
                # Extract image path
                try:
                    img_path = line.split('](')[1].split(')')[0]
                    img_file = output_dir / img_path
                    if img_file.exists() and img_file.is_file():
                        doc.add_picture(str(img_file), width=Inches(6))
                        doc.add_paragraph()  # Add empty line
                except Exception:
                    pass  # If image doesn't exist, skip
                i += 1
                continue
            
            # Heading processing
            if line.startswith('# '):
                # Level 1 heading
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                # Level 2 heading
                p = doc.add_heading(line[3:], level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                # Level 3 heading
                p = doc.add_heading(line[4:], level=3)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('---'):
                # Separator line
                doc.add_paragraph('─' * 50)
            elif line.startswith('|'):
                # Table
                table_lines = []
                while i < len(md_lines) and md_lines[i].strip().startswith('|'):
                    table_lines.append(md_lines[i].strip())
                    i += 1
                i -= 1  # Step back one
        
                if len(table_lines) >= 2:
                    # Parse table header
                    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                    # Skip separator row
                    data_rows = table_lines[2:]
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add table header
                    header_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        header_cells[j].text = header
                        header_cells[j].paragraphs[0].runs[0].font.bold = True
                    
                    # Add data rows
                    for row_line in data_rows:
                        if row_line.strip().startswith('| ...'):
                            continue  # Skip ellipsis row
                        cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        if len(cells) == len(headers):
                            row_cells = table.add_row().cells
                            for j, cell_text in enumerate(cells):
                                row_cells[j].text = cell_text
            elif line.startswith('- ') or line.startswith('* '):
                # List item
                text = line[2:].strip()
                # Process bold
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                if text:  # Ensure text is not empty
                    doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\d+\.\s+', line):
                # Ordered list (match any number-prefixed list item)
                text = re.sub(r'^\d+\.\s+', '', line)
                # Process bold
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                if text:  # Ensure text is not empty
                    doc.add_paragraph(text, style='List Number')
            else:
                # Normal paragraph
                # Process bold
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                if text:  # Ensure text is not empty
                    doc.add_paragraph(text)
            
            i += 1
        
        # Save Word document
        word_file = output_dir / "detailed_analysis_report.docx"
        doc.save(str(word_file))
        print(f"✓ Detailed analysis report (Word) saved: {word_file}")
    except Exception as e:
        print(f"  ⚠ Word report generation failed: {e}")


def _generate_pdf_report(md_lines: list[str], output_dir: Path) -> None:
    """Convert Markdown content to PDF document"""
    try:
        import markdown
    except ImportError:
        print("  ⚠ markdown not installed, skipping PDF report generation")
        print("    Tip: Run pip install markdown to generate PDF format reports")
        print("    Or: Use Word document and save as PDF in Microsoft Word")
        return
    
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
    except (ImportError, OSError) as e:
        print("  ⚠ weasyprint not installed or system dependencies missing, skipping PDF report generation")
        print(f"    Error: {e}")
        print("    Tip: weasyprint requires additional system libraries on Windows, installation is complex")
        print("    Suggestion: Use Word document and save as PDF in Microsoft Word")
        return
    
    try:
        # Convert Markdown to HTML
        md_content = "\n".join(md_lines)
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code']
        )
        
        # Add styles
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    font-size: 24pt;
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    margin-top: 30px;
                }}
                h2 {{
                    font-size: 18pt;
                    color: #34495e;
                    margin-top: 25px;
                    border-bottom: 2px solid #ecf0f1;
                    padding-bottom: 5px;
                }}
                h3 {{
                    font-size: 14pt;
                    color: #7f8c8d;
                    margin-top: 20px;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                    font-size: 10pt;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #3498db;
                    color: white;
                    font-weight: bold;
                }}
                tr:nth-child(even) {{
                    background-color: #f9f9f9;
                }}
                ul, ol {{
                    margin: 10px 0;
                    padding-left: 30px;
                }}
                li {{
                    margin: 5px 0;
                }}
                p {{
                    margin: 10px 0;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        pdf_file = output_dir / "detailed_analysis_report.pdf"
        font_config = FontConfiguration()
        HTML(string=styled_html).write_pdf(
            pdf_file,
            font_config=font_config
        )
        print(f"✓ Detailed analysis report (PDF) saved: {pdf_file}")
    except (OSError, Exception) as e:
        print(f"  ⚠ PDF report generation failed: {e}")
        print("    Tip: weasyprint may require additional system libraries on Windows")
        print("    Suggestion: Use Word document and save as PDF in Microsoft Word")


def create_summary_excel(summary: dict[str, Any], output_dir: Path, benchmark: Any = None) -> None:
    """
    Create summary Excel report, including comparison of all models and statistics by category/difficulty.
    
    Args:
        summary: Evaluation summary result dictionary, containing fields like results, benchmark, total_questions
        output_dir: Output directory path
        benchmark: Benchmark object (optional, currently unused but kept for interface compatibility)
    """
    if pd is None:
        print("⚠ pandas not installed, skipping summary Excel report generation")
        return
    
    # Model comparison table
    rows = []
    for result in summary["results"]:
        rows.append({
            "Model Name": result["model_name"],
            "Correct Count": result["total_score"],
            "Total": result["max_score"],
            "Accuracy": result["accuracy"],  # Use numeric value, Excel will automatically format as percentage
            "Incorrect": result["incorrect_count"],
            "Missing": result["missing_count"],
        })
    
    df = pd.DataFrame(rows)
    
    summary_file = output_dir / "models_comparison.xlsx"
    with pd.ExcelWriter(summary_file, engine="openpyxl") as writer:
        # Model comparison worksheet
        df.to_excel(writer, sheet_name="Model Comparison", index=False)
        worksheet = writer.sheets["Model Comparison"]
        worksheet.column_dimensions["A"].width = 40  # Model Name
        worksheet.column_dimensions["B"].width = 12  # Correct Count
        worksheet.column_dimensions["C"].width = 12  # Total
        worksheet.column_dimensions["D"].width = 12  # Accuracy
        worksheet.column_dimensions["E"].width = 10  # Incorrect
        worksheet.column_dimensions["F"].width = 10  # Missing
        
        # Set accuracy column to percentage format
        from openpyxl.styles import numbers
        for row in range(2, len(df) + 2):  # Start from row 2 (row 1 is header)
            cell = worksheet.cell(row, 4)  # Column D is accuracy
            cell.number_format = numbers.FORMAT_PERCENTAGE_00  # Percentage format, 2 decimal places
        
        # Add summary information
        summary_row = len(df) + 3
        worksheet.cell(summary_row, 1, "Benchmark Info")
        worksheet.cell(summary_row + 1, 1, "Benchmark: hydrosebench")
        worksheet.cell(summary_row + 2, 1, f"Total Questions: {summary['total_questions']}")
        worksheet.cell(summary_row + 3, 1, f"Models Evaluated: {summary['models_count']}")
        
        # Statistics by category (comparison of all models)
        if any(result.get("category_stats") for result in summary["results"]):
            # Collect all categories
            all_categories = set()
            for result in summary["results"]:
                if result.get("category_stats"):
                    all_categories.update(result["category_stats"].keys())
            
            if all_categories:
                category_rows = []
                for category in sorted(all_categories):
                    row = {"Category": category}
                    for result in summary["results"]:
                        model_name = result["model_name"]
                        stats = result.get("category_stats", {}).get(category, {})
                        if stats:
                            # Separate into two columns: score and accuracy
                            row[f"{model_name}_Score"] = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                            row[f"{model_name}_Accuracy"] = stats.get('accuracy', 0)  # Numeric format
                        else:
                            row[f"{model_name}_Score"] = "-"
                            row[f"{model_name}_Accuracy"] = None
                    category_rows.append(row)
                
                category_df = pd.DataFrame(category_rows)
                category_df.to_excel(writer, sheet_name="By Category Comparison", index=False)
                cat_worksheet = writer.sheets["By Category Comparison"]
                cat_worksheet.column_dimensions["A"].width = 30  # Category
                
                # Set column widths and formats
                col_idx = 2
                for result in summary["results"]:
                    model_name = result["model_name"]
                    # Score column
                    cat_worksheet.column_dimensions[chr(64 + col_idx)].width = 20
                    col_idx += 1
                    # Accuracy column - set to percentage format
                    from openpyxl.styles import numbers
                    for row in range(2, len(category_df) + 2):
                        cell = cat_worksheet.cell(row, col_idx)
                        if cell.value is not None:
                            cell.number_format = numbers.FORMAT_PERCENTAGE_00
                    cat_worksheet.column_dimensions[chr(64 + col_idx)].width = 15
                    col_idx += 1
        
        # Statistics by difficulty level (comparison of all models)
        if any(result.get("level_stats") for result in summary["results"]):
            # Collect all difficulty levels
            all_levels = set()
            for result in summary["results"]:
                if result.get("level_stats"):
                    all_levels.update(result["level_stats"].keys())
            
            if all_levels:
                level_rows = []
                for level in sorted(all_levels):
                    row = {"Level": level}
                    for result in summary["results"]:
                        model_name = result["model_name"]
                        stats = result.get("level_stats", {}).get(level, {})
                        if stats:
                            # Separate into two columns: score and accuracy
                            row[f"{model_name}_Score"] = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                            row[f"{model_name}_Accuracy"] = stats.get('accuracy', 0)  # Numeric format
                        else:
                            row[f"{model_name}_Score"] = "-"
                            row[f"{model_name}_Accuracy"] = None
                    level_rows.append(row)
                
                level_df = pd.DataFrame(level_rows)
                level_df.to_excel(writer, sheet_name="By Level Comparison", index=False)
                level_worksheet = writer.sheets["By Level Comparison"]
                level_worksheet.column_dimensions["A"].width = 30  # Level
                
                # Set column widths and formats
                col_idx = 2
                for result in summary["results"]:
                    model_name = result["model_name"]
                    # Score column
                    level_worksheet.column_dimensions[chr(64 + col_idx)].width = 20
                    col_idx += 1
                    # Accuracy column - set to percentage format
                    from openpyxl.styles import numbers
                    for row in range(2, len(level_df) + 2):
                        cell = level_worksheet.cell(row, col_idx)
                        if cell.value is not None:
                            cell.number_format = numbers.FORMAT_PERCENTAGE_00
                    level_worksheet.column_dimensions[chr(64 + col_idx)].width = 15
                    col_idx += 1
        
        # Statistics by question type (comparison of all models)
        if any(result.get("type_stats") for result in summary["results"]):
            # Collect all question types
            all_types = set()
            for result in summary["results"]:
                if result.get("type_stats"):
                    all_types.update(result["type_stats"].keys())
            
            if all_types:
                type_rows = []
                for qtype in sorted(all_types):
                    row = {"Type": qtype}
                    for result in summary["results"]:
                        model_name = result["model_name"]
                        stats = result.get("type_stats", {}).get(qtype, {})
                        if stats:
                            # Separate into two columns: score and accuracy
                            row[f"{model_name}_Score"] = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                            row[f"{model_name}_Accuracy"] = stats.get('accuracy', 0)  # Numeric format
                        else:
                            row[f"{model_name}_Score"] = "-"
                            row[f"{model_name}_Accuracy"] = None
                    type_rows.append(row)
                
                type_df = pd.DataFrame(type_rows)
                type_df.to_excel(writer, sheet_name="By Type Comparison", index=False)
                type_worksheet = writer.sheets["By Type Comparison"]
                type_worksheet.column_dimensions["A"].width = 30  # Type
                
                # Set column widths and formats
                col_idx = 2
                for result in summary["results"]:
                    model_name = result["model_name"]
                    # Score column
                    type_worksheet.column_dimensions[chr(64 + col_idx)].width = 20
                    col_idx += 1
                    # Accuracy column - set to percentage format
                    from openpyxl.styles import numbers
                    for row in range(2, len(type_df) + 2):
                        cell = type_worksheet.cell(row, col_idx)
                        if cell.value is not None:
                            cell.number_format = numbers.FORMAT_PERCENTAGE_00
                    type_worksheet.column_dimensions[chr(64 + col_idx)].width = 15
                    col_idx += 1
    
    print(f"✓ Model comparison summary saved: {summary_file}")


def validate_data_quality(df: pd.DataFrame, benchmark: Benchmark) -> dict[str, Any]:
    """
    Validate data quality.
    
    Args:
        df: DataFrame of prediction data
        benchmark: Benchmark object
    
    Returns:
        Dictionary containing data quality check results
    """
    if pd is None:
        raise ImportError("pandas not installed, cannot perform data quality check")
    
    issues = []
    warnings = []
    
    # Check ID column
    if "ID" not in df.columns:
        issues.append("Missing ID column")
    else:
        # Check ID matching
        df_ids = set(df["ID"].astype(str))
        benchmark_ids = set(ex.id for ex in benchmark.examples)
        missing_ids = benchmark_ids - df_ids
        extra_ids = df_ids - benchmark_ids
        
        if missing_ids:
            warnings.append(f"Benchmark has {len(missing_ids)} IDs missing in prediction file")
        if extra_ids:
            warnings.append(f"Prediction file has {len(extra_ids)} IDs not in benchmark")
    
    # Check null values
    model_cols = identify_model_columns(df, verbose=False)
    for col in model_cols:
        null_count = df[col].isna().sum()
        if null_count > 0:
            warnings.append(f"Model column '{col}' has {null_count} null values")
    
    return {
        "issues": issues,
        "warnings": warnings,
        "model_count": len(model_cols),
        "data_rows": len(df),
        "benchmark_size": len(benchmark.examples),
    }


def analyze_errors(summary: dict[str, Any], benchmark: Benchmark, output_dir: Path) -> dict[str, Any]:
    """
    Deep error analysis.
    
    Args:
        summary: Evaluation summary results
        benchmark: Benchmark object
        output_dir: Output directory, used to read detailed reports
    
    Returns:
        Error analysis result dictionary
    """
    error_analysis = {
        "common_errors": {},  # question_id -> list of error models
        "hard_questions": [],  # Questions that all models got wrong
        "easy_questions": [],  # Questions that all models got correct
        "model_errors": {},  # model -> list of error questions
    }
    
    # Collect error situation for each question
    question_errors = {}  # question_id -> {correct_models: [], incorrect_models: []}
    
    for result in summary["results"]:
        model_name = result["model_name"]
        error_analysis["model_errors"][model_name] = []
        
        # Read detailed report for this model
        # Generate safe folder name
        safe_name = model_name.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?", "_").replace('"', "_").replace("<", "_").replace(">", "_").replace("|", "_")
        model_dir = output_dir / safe_name
        report_file = model_dir / "score_report.json"
        if report_file.exists():
            with open(report_file, "r", encoding="utf-8") as f:
                report_data = json.load(f)
            
            for example in report_data.get("examples", []):
                qid = example["id"]
                is_correct = example["is_correct"]
                
                if qid not in question_errors:
                    question_errors[qid] = {"correct": [], "incorrect": []}
                
                if is_correct:
                    question_errors[qid]["correct"].append(model_name)
                else:
                    question_errors[qid]["incorrect"].append(model_name)
                    error_analysis["model_errors"][model_name].append(qid)
    
    # Analyze question difficulty
    total_models = len(summary["results"])
    for qid, stats in question_errors.items():
        correct_count = len(stats["correct"])
        incorrect_count = len(stats["incorrect"])
        
        if incorrect_count == total_models:
            error_analysis["hard_questions"].append(qid)
        elif correct_count == total_models:
            error_analysis["easy_questions"].append(qid)
        
        if incorrect_count > total_models * 0.5:  # More than half of models got it wrong
            error_analysis["common_errors"][qid] = {
                "error_rate": incorrect_count / total_models,
                "incorrect_models": stats["incorrect"],
            }
    
    return error_analysis


def generate_analysis_report(
    summary: dict[str, Any],
    benchmark: Benchmark,
    output_dir: Path,
) -> dict[str, Any]:
    """
    Generate detailed analysis report (Markdown format).
    
    Args:
        summary: Evaluation summary results
        benchmark: Benchmark object
        output_dir: Output directory
    
    Returns:
        Analysis report dictionary (JSON format)
    """
    if pd is None:
        raise ImportError("pandas not installed, cannot generate analysis report")
    
    best_model = summary["results"][0]
    worst_model = summary["results"][-1]
    avg_accuracy = sum(r["accuracy"] for r in summary["results"]) / len(summary["results"])
    std_accuracy = pd.Series([r["accuracy"] for r in summary["results"]]).std()
    accuracy_gap = best_model["accuracy"] - worst_model["accuracy"]
    
    # Generate Markdown report
    md_lines = []
    md_lines.append("# Detailed Analysis Report")
    md_lines.append("")
    md_lines.append(f"**Generated Time**: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    
    # 1. Evaluation Overview
    md_lines.append("## 📊 Evaluation Overview")
    md_lines.append("")
    md_lines.append(f"- **Number of Models Evaluated**: {summary['models_count']}")
    md_lines.append(f"- **Total Questions**: {summary['total_questions']}")
    md_lines.append(f"- **Average Accuracy**: {avg_accuracy:.2%}")
    md_lines.append(f"- **Accuracy Standard Deviation**: {std_accuracy:.4f}")
    md_lines.append("")
    
    # 2. Model Rankings
    md_lines.append("## 🏆 Model Rankings")
    md_lines.append("")
    md_lines.append("| Rank | Model Name | Score | Accuracy |")
    md_lines.append("|------|-----------|-------|----------|")
    for idx, result in enumerate(summary["results"], 1):
        score_str = f"{result['total_score']}/{result['max_score']}"
        accuracy_str = f"{result['accuracy']:.2%}"
        md_lines.append(f"| {idx} | {result['model_name']} | {score_str} | {accuracy_str} |")
    md_lines.append("")
    
    # 3. Key Metrics
    md_lines.append("## 📈 Key Metrics")
    md_lines.append("")
    md_lines.append(f"### Best Model")
    md_lines.append(f"- **Model**: {best_model['model_name']}")
    md_lines.append(f"- **Accuracy**: {best_model['accuracy']:.2%}")
    md_lines.append(f"- **Score**: {best_model['total_score']}/{best_model['max_score']}")
    md_lines.append("")
    md_lines.append(f"### Worst Model")
    md_lines.append(f"- **Model**: {worst_model['model_name']}")
    md_lines.append(f"- **Accuracy**: {worst_model['accuracy']:.2%}")
    md_lines.append(f"- **Score**: {worst_model['total_score']}/{worst_model['max_score']}")
    md_lines.append("")
    md_lines.append(f"### Model Gap")
    md_lines.append(f"- **Accuracy Gap**: {accuracy_gap:.2%}")
    md_lines.append("")
    
    # 4. Category Analysis
    if summary["results"] and summary["results"][0].get("category_stats"):
        md_lines.append("## 📂 Category Dimension Analysis")
        md_lines.append("")
        all_categories = set()
        for result in summary["results"]:
            if result.get("category_stats"):
                all_categories.update(result["category_stats"].keys())
        
        for category in sorted(all_categories):
            md_lines.append(f"### {category}")
            md_lines.append("")
            md_lines.append("| Model Name | Accuracy | Score |")
            md_lines.append("|-----------|----------|-------|")
            
            best_acc = 0
            best_model_name = ""
            worst_acc = 1.0
            worst_model_name = ""
            
            category_results = []
            for result in summary["results"]:
                stats = result.get("category_stats", {}).get(category, {})
                if stats:
                    acc = stats.get("accuracy", 0)
                    score_str = f"{stats.get('correct', 0)}/{stats.get('total', 0)}"
                    category_results.append((result["model_name"], acc, score_str))
                    if acc > best_acc:
                        best_acc = acc
                        best_model_name = result["model_name"]
                    if acc < worst_acc:
                        worst_acc = acc
                        worst_model_name = result["model_name"]
            
            # Sort by accuracy
            category_results.sort(key=lambda x: x[1], reverse=True)
            for model_name, acc, score_str in category_results:
                md_lines.append(f"| {model_name} | {acc:.2%} | {score_str} |")
            
            md_lines.append("")
            md_lines.append(f"- **Best**: {best_model_name} ({best_acc:.2%})")
            md_lines.append(f"- **Worst**: {worst_model_name} ({worst_acc:.2%})")
            md_lines.append(f"- **Gap**: {(best_acc - worst_acc):.2%}")
            md_lines.append("")
    
    # 5. Difficulty Analysis
    if any(r.get("level_stats") for r in summary["results"]):
        md_lines.append("## 🎯 Difficulty Dimension Analysis")
        md_lines.append("")
        all_levels = set()
        for result in summary["results"]:
            if result.get("level_stats"):
                all_levels.update(result["level_stats"].keys())
        
        md_lines.append("| Difficulty Level | Average Accuracy | Number of Models |")
        md_lines.append("|----------------|------------------|------------------|")
        
        for level in sorted(all_levels):
            level_accuracies = []
            for result in summary["results"]:
                stats = result.get("level_stats", {}).get(level, {})
                if stats:
                    level_accuracies.append(stats.get("accuracy", 0))
            
            if level_accuracies:
                avg_level_acc = sum(level_accuracies) / len(level_accuracies)
                md_lines.append(f"| {level} | {avg_level_acc:.2%} | {len(level_accuracies)} |")
        
        md_lines.append("")
    
    # 6. Recommendations
    md_lines.append("## 💡 Key Recommendations")
    md_lines.append("")
    recommendations = []
    recommendations.append(f"✅ Best model **{best_model['model_name']}** has accuracy of {best_model['accuracy']:.2%}, recommend as baseline model")
    recommendations.append(f"⚠️ Accuracy gap between models is {accuracy_gap:.2%}, recommend analyzing performance difference reasons")
    
    if avg_accuracy < 0.5:
        recommendations.append("⚠️ Average accuracy below 50%, recommend checking model configuration or data quality")
    elif avg_accuracy > 0.9:
        recommendations.append("✅ Average accuracy is high, recommend increasing question difficulty or expanding benchmark scale")
    
    for i, rec in enumerate(recommendations, 1):
        md_lines.append(f"{i}. {rec}")
    md_lines.append("")
    
    # 8. Detailed Data (link to JSON)
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 📎 Related Files")
    md_lines.append("")
    md_lines.append("- Complete data (JSON format): `detailed_analysis_report.json`")
    md_lines.append("- Model comparison summary: `models_comparison.xlsx`")
    md_lines.append("- Detailed reports for each model: `<model_name>/score_report.xlsx` and `<model_name>/score_report.json`")
    md_lines.append("- Word format report: `detailed_analysis_report.docx` (if python-docx is installed)")
    md_lines.append("- PDF format report: `detailed_analysis_report.pdf` (if weasyprint is installed)")
    md_lines.append("")
    
    # Save Markdown report
    md_file = output_dir / "detailed_analysis_report.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    
    print(f"✓ Detailed analysis report (Markdown) saved: {md_file}")
    
    # Try to generate Word and PDF formats
    _generate_word_report(md_lines, output_dir)
    _generate_pdf_report(md_lines, output_dir)
    
    # Also save JSON format as data backup
    report_json = {
        "summary": {
            "total_models": summary["models_count"],
            "total_questions": summary["total_questions"],
            "evaluation_date": pd.Timestamp.now().isoformat(),
            "average_accuracy": avg_accuracy,
            "std_accuracy": float(std_accuracy),
        },
        "model_rankings": [
            {
                "rank": idx,
                "model_name": result["model_name"],
                "accuracy": result["accuracy"],
                "score": f"{result['total_score']}/{result['max_score']}",
            }
            for idx, result in enumerate(summary["results"], 1)
        ],
        "key_metrics": {
            "best_model": {
                "name": best_model["model_name"],
                "accuracy": best_model["accuracy"],
                "score": f"{best_model['total_score']}/{best_model['max_score']}",
            },
            "worst_model": {
                "name": worst_model["model_name"],
                "accuracy": worst_model["accuracy"],
                "score": f"{worst_model['total_score']}/{worst_model['max_score']}",
            },
            "accuracy_gap": accuracy_gap,
        },
        "recommendations": recommendations,
    }
    
    json_file = output_dir / "detailed_analysis_report.json"
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(report_json, f, ensure_ascii=False, indent=2)
    
    print(f"✓ Detailed analysis report (JSON data) saved: {json_file}")
    
    return report_json

